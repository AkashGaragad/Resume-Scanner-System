# app_advanced.py - Advanced web application with proper scoring
import streamlit as st
import json
import pdfplumber
import io
from docx import Document
from datetime import datetime, timedelta
from scoring_engine import calculate_hard_match_score, calculate_semantic_similarity, generate_gemini_feedback
from database import (
    init_database, save_evaluation, get_all_evaluations, 
    search_evaluations, delete_evaluation, delete_all_evaluations, 
    cleanup_old_records
)

# Helper function to safely load JSON data from the database
# This remains useful as data is returned from the DB as a string representation of JSON
def safe_json_loads(json_data, default_value=None):
    """Safely loads a JSON object (or string), returning a default value if it fails."""
    if default_value is None:
        default_value = []
    
    if not json_data:
        return default_value
        
    # If it's already a list or dict, just return it
    if isinstance(json_data, (list, dict)):
        return json_data
        
    try:
        # If it's a string, try to load it
        return json.loads(json_data)
    except (json.JSONDecodeError, TypeError):
        return default_value

# --- Initialization ---
# This will now initialize the PostgreSQL table
init_database()
 # Run this once on startup to fix any existing bad data

# --- Page Configuration ---
st.set_page_config(page_title="Advanced Resume Scanner", page_icon="📊", layout="wide")

# --- Page Title and Description ---
st.title("📊 Advanced Resume Relevance Checker")
st.write("Professional-grade resume analysis using weighted scoring and semantic similarity.")

# --- File Handling ---
def read_file(file):
    """Reads text content from uploaded PDF, DOCX, or TXT files."""
    text = ""
    file_type = file.type
    
    try:
        if file_type == "application/pdf":
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            docx_file = io.BytesIO(file.read())
            doc = Document(docx_file)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
        else: # Assumes .txt
            text = str(file.read(), "utf-8")
            
    except Exception as e:
        st.error(f"Error reading file '{file.name}': {str(e)}")
        return ""
    
    return text

# --- Load Skills Data ---
@st.cache_data
def load_skills_data():
    with open('skills.json', 'r') as f:
        return json.load(f)

skills_library = load_skills_data()

# --- Sidebar ---
with st.sidebar:
    st.header("Upload Files")
    # Added 'docx' to supported file types
    uploaded_jd = st.file_uploader("Job Description (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'])
    uploaded_resume = st.file_uploader("Resume (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'])
    
    st.divider()
    st.header("History Management")
    
    days_to_keep = st.number_input("Delete records older than (days):", min_value=1, value=30)
    if st.button("🧹 Clean Up Old Records", type="secondary"):
        deleted_count = cleanup_old_records(days_to_keep)
        st.success(f"Deleted {deleted_count} records older than {days_to_keep} days!")
        st.rerun()

    if st.button("🗑️ Delete All History", type="secondary"):
        if 'confirm_delete' not in st.session_state:
            st.session_state.confirm_delete = True
        else:
            del st.session_state.confirm_delete
        st.rerun()

    if 'confirm_delete' in st.session_state and st.session_state.confirm_delete:
        st.warning("This will delete ALL evaluation history. This action cannot be undone!")
        col1, col2 = st.columns(2)
        if col1.button("✅ Confirm Delete All", type="primary"):
            delete_all_evaluations()
            st.success("All history has been deleted!")
            del st.session_state.confirm_delete
            st.rerun()
        if col2.button("Cancel"):
            del st.session_state.confirm_delete
            st.rerun()

    st.divider()
    st.info("""
    **Scoring Methodology:**
    - 45% Hard Match (Skills, Education, Concepts, Experience)
    - 55% Semantic Similarity (Contextual understanding)
    """)

# --- Main Page Tabs ---
tab1, tab2 = st.tabs(["🧪 New Evaluation", "📊 Evaluation History"])

with tab1:
    st.header("Run New Analysis")
    if st.button("Run Advanced Analysis 🚀", type="primary"):
        if uploaded_jd and uploaded_resume:
            with st.spinner('Running advanced analysis... This may take a moment.'):
                jd_text = read_file(uploaded_jd)
                resume_text = read_file(uploaded_resume)

                if jd_text and resume_text:
                    # --- Scoring and Analysis ---
                    hard_score, score_breakdown = calculate_hard_match_score(jd_text, resume_text, skills_library)
                    semantic_score = calculate_semantic_similarity(jd_text, resume_text)
                    final_score = (hard_score * 0.45) + (semantic_score * 0.55)

                    # --- AI Feedback Generation ---
                    with st.spinner('🧠 Generating personalized AI feedback...'):
                        all_found = score_breakdown['found_must_have'].union(score_breakdown['found_good_to_have'])
                        all_required = score_breakdown['must_have_skills'].union(score_breakdown['good_to_have_skills'])
                        all_missing = all_required - all_found
                        
                        api_key = st.secrets.get("GEMINI_API_KEY")
                        if api_key:
                            ai_feedback = generate_gemini_feedback(
                                jd_text, resume_text, all_missing, final_score, all_found, api_key=api_key
                            )
                        else:
                            ai_feedback = "⚠️ Gemini API key not found. Please add it to your Streamlit secrets."
                    
                    # --- Save to Database ---
                    result_data = {
                        'jd_filename': uploaded_jd.name, 'resume_filename': uploaded_resume.name,
                        'hard_score': hard_score, 'semantic_score': semantic_score, 'final_score': final_score,
                        'verdict': 'HIGH FIT' if final_score >= 70 else 'MEDIUM FIT' if final_score >= 40 else 'LOW FIT',
                        'must_have_skills': score_breakdown['must_have_skills'],
                        'found_must_have_skills': score_breakdown['found_must_have'],
                        'good_to_have_skills': score_breakdown['good_to_have_skills'],
                        'found_good_to_have_skills': score_breakdown['found_good_to_have'],
                        'ai_feedback': ai_feedback
                    }
                    save_evaluation(result_data)

                    # --- Display Results ---
                    st.header("📈 Advanced Analysis Results")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Hard Match Score", f"{hard_score:.1f}%")
                    col2.metric("Semantic Score", f"{semantic_score:.1f}%")
                    col3.metric("Final Score", f"{final_score:.1f}%", delta=result_data['verdict'])
                    st.progress(final_score / 100)
                    
                    with st.expander("📋 Detailed Score Breakdown"):
                        st.subheader("Hard Match Components")
                        c1, c2 = st.columns(2)
                        c1.write(f"**Skills Score:** {score_breakdown['skill_score']:.1f}%")
                        c1.write(f"**Education Score:** {score_breakdown['education_score']:.1f}%")
                        c2.write(f"**Concept Match Score:** {score_breakdown['concept_score']:.1f}%")
                        c2.write(f"**Experience Score:** {score_breakdown['experience_score']:.1f}%")

                    st.subheader("Skills Analysis")
                    c1, c2 = st.columns(2)
                    
                    must_have = score_breakdown.get('must_have_skills', set())
                    found_must = score_breakdown.get('found_must_have', set())
                    good_to_have = score_breakdown.get('good_to_have_skills', set())
                    found_good = score_breakdown.get('found_good_to_have', set())
                    
                    with c1:
                        st.write("**Must-Have Skills:**")
                        if must_have:
                            for skill in must_have:
                                st.write(f"{'✅' if skill in found_must else '❌'} {skill}")
                        else:
                            st.write("None specified in JD.")
                            
                    with c2:
                        st.write("**Good-to-Have Skills:**")
                        if good_to_have:
                            for skill in good_to_have:
                                st.write(f"{'✅' if skill in found_good else '➖'} {skill}")
                        else:
                            st.write("None specified in JD.")
                    
                    st.subheader("🎯 Final Verdict")
                    if final_score >= 70:
                        st.success("**HIGH FIT** - Strong alignment with job requirements.")
                    elif final_score >= 40:
                        st.warning("**MEDIUM FIT** - Partial alignment with requirements.")
                    else:
                        st.error("**LOW FIT** - Limited alignment with requirements.")

                    with st.expander("🤖 AI Career Coach Feedback", expanded=True):
                        st.markdown(ai_feedback)
        else:
            st.warning("Please upload both a Job Description and a Resume to begin.")

with tab2:
    st.header("📜 Evaluation History")
    
    c1, c2 = st.columns([3, 1])
    search_term = c1.text_input("🔍 Search by filename or verdict:")
    verdict_filter = c2.selectbox("Filter by verdict:", ["All", "HIGH FIT", "MEDIUM FIT", "LOW FIT"])
    
    evaluations = search_evaluations(search_term) if search_term else get_all_evaluations()
    
    if verdict_filter != "All":
        evaluations = [e for e in evaluations if e['verdict'] == verdict_filter]
    
    if evaluations:
        st.write(f"**Found {len(evaluations)} evaluations**")
        
        for eval_item in evaluations:
            timestamp_str = eval_item['timestamp'].strftime('%Y-%m-%d %H:%M')
            expander_title = (
                f"📋 {eval_item['resume_filename']} vs {eval_item['job_description_filename']} | "
                f"{eval_item['verdict']} | Score: {eval_item['final_score']:.1f}%"
            )
            with st.expander(expander_title):
                c1, c2 = st.columns(2)
                c1.write(f"**📅 Date:** {timestamp_str}")
                c1.write(f"**🎯 Hard Score:** {eval_item['hard_score']:.1f}%")
                c2.write(f"**🧠 Semantic Score:** {eval_item['semantic_score']:.1f}%")
                c2.write(f"**✅ Final Score:** {eval_item['final_score']:.1f}%")
                
                st.subheader("Skills Analysis")
                c3, c4 = st.columns(2)
                
                must_have = safe_json_loads(eval_item['must_have_skills'])
                found_must = safe_json_loads(eval_item['found_must_have_skills'])
                missing_must = list(set(must_have) - set(found_must))
                
                with c3:
                    st.write("**Must-Have Skills Found:**")
                    if found_must:
                        for skill in found_must:
                            st.write(f"✅ {skill}")
                    else:
                        st.write("None")

                with c4:
                    st.write("**Must-Have Skills Missing:**")
                    if missing_must:
                        for skill in missing_must:
                            st.write(f"❌ {skill}")
                    elif must_have:
                        st.write("All skills found! 👍")
                    else:
                        st.write("No 'must-have' skills defined.")
                
                if eval_item['ai_feedback']:
                    with st.expander("🤖 AI Feedback"):
                        st.markdown(eval_item['ai_feedback'])
                
                if st.button("🗑️ Delete This Entry", key=f"delete_{eval_item['id']}"):
                    delete_evaluation(eval_item['id'])
                    st.success(f"Evaluation for '{eval_item['resume_filename']}' deleted!")
                    st.rerun()
    else:
        st.info("No evaluations found. Run a new analysis or clear your search filters.")
